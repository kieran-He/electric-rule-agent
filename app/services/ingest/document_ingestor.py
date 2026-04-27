import hashlib
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

from app.config import settings
from app.core.repository import ChromaPolicyRepository


SUPPORTED_SUFFIXES = {".txt", ".pdf", ".docx"}


@dataclass
class IngestStats:
    files_processed: int = 0
    chunks_created: int = 0
    files_new: int = 0
    files_updated: int = 0
    files_skipped: int = 0
    ocr_pages_processed: int = 0


class FingerprintStore:
    def __init__(self, index_path: str):
        self.path = Path(index_path)
        self.path.parent.mkdir(parents=True, exist_ok=True)

    def load(self) -> Dict[str, str]:
        if not self.path.exists():
            return {}
        try:
            return json.loads(self.path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            return {}

    def save(self, data: Dict[str, str]) -> None:
        self.path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    def clear_scope(self, data: Dict[str, str], kb_scope: str, province_code: Optional[str]) -> Dict[str, str]:
        scope_key = f"{kb_scope}:{(province_code or '').upper()}:"
        return {k: v for k, v in data.items() if not k.startswith(scope_key)}


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _strip_control(text: str) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)


def _extract_effective_date(text: str) -> str:
    patterns = [
        r"(20\d{2})年\s*(\d{1,2})月\s*(\d{1,2})日",
        r"(20\d{2})[\-/\.](\d{1,2})[\-/\.](\d{1,2})",
        r"(20\d{2})年\s*(\d{1,2})月",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            parts = [x.zfill(2) for x in m.groups()]
            if len(parts) == 3:
                return f"{parts[0]}-{parts[1]}-{parts[2]}"
            return f"{parts[0]}-{parts[1]}-01"
    return ""


def _extract_policy_level(text: str) -> str:
    sample = text[:800]
    if any(k in sample for k in ["国家发展改革委", "国家能源局", "中华人民共和国"]):
        return "national"
    if "省" in sample or "自治区" in sample or "直辖市" in sample:
        return "province"
    return "unknown"


def _text_quality_metrics(text: str) -> Dict[str, float]:
    if not text:
        return {"ch_ratio": 0.0, "replacement_ratio": 1.0, "readable_density": 0.0}
    visible_chars = [c for c in text if not c.isspace()]
    visible = len(visible_chars)
    if visible == 0:
        return {"ch_ratio": 0.0, "replacement_ratio": 1.0, "readable_density": 0.0}
    chinese = len(re.findall(r"[\u4e00-\u9fff]", text))
    replacement = text.count("")
    readable = len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9\uFF01\uFF0C\uFF1F\uFF1B\uFF1A\u201C\u201D\u2018\u2019\uFF08\uFF09\u3010\u3011\u300A\u300B,.!?;:\-]", text))
    return {
        "ch_ratio": chinese / visible,
        "replacement_ratio": replacement / visible,
        "readable_density": readable / max(1, len(text)),
    }


def _is_low_quality(text: str, min_ch_ratio: float, max_replacement_ratio: float) -> bool:
    metrics = _text_quality_metrics(text)
    return metrics["ch_ratio"] < min_ch_ratio or metrics["replacement_ratio"] > max_replacement_ratio


def _clean_lines(lines: List[str]) -> List[str]:
    cleaned: List[str] = []
    for raw in lines:
        line = _strip_control(raw).strip()
        if not line:
            continue
        if re.fullmatch(r"[-—_=\s\d]{1,16}", line):
            continue
        if "目录" in line and ("..." in line or "……" in line):
            continue
        if re.search(r"\.{3,}|…{2,}", line):
            continue
        cleaned.append(line)
    return cleaned


def _drop_high_frequency_lines(page_lines: List[List[str]]) -> List[List[str]]:
    if not page_lines:
        return page_lines
    total_pages = len(page_lines)
    freq: Dict[str, int] = {}
    for lines in page_lines:
        for line in set(lines):
            if len(line) <= 60:
                freq[line] = freq.get(line, 0) + 1
    noisy = {k for k, v in freq.items() if v >= max(3, int(total_pages * 0.6))}
    return [[ln for ln in lines if ln not in noisy] for lines in page_lines]


def _normalize_from_pages(pages: List[str]) -> str:
    split_pages = [_clean_lines((p or "").splitlines()) for p in pages]
    split_pages = _drop_high_frequency_lines(split_pages)
    merged = ["\n".join(lines) for lines in split_pages if lines]
    text = "\n\n".join(merged)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_text(text: str, chunk_size: int = 800, chunk_overlap: int = 120) -> List[str]:
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        if len(para) > chunk_size:
            if current:
                chunks.append(current)
                current = ""
            start = 0
            while start < len(para):
                end = min(start + chunk_size, len(para))
                chunks.append(para[start:end])
                if end == len(para):
                    break
                start = max(end - chunk_overlap, start + 1)
            continue
        if not current:
            current = para
            continue
        candidate = f"{current}\n\n{para}"
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            chunks.append(current)
            overlap_text = current[-chunk_overlap:] if chunk_overlap > 0 else ""
            current = (overlap_text + "\n\n" + para).strip()
            if len(current) > chunk_size:
                chunks.append(current[:chunk_size])
                current = current[chunk_size - chunk_overlap :]
    if current:
        chunks.append(current)
    return [c for c in chunks if len(c) >= 40]


def _extract_pdf_pages(path: Path) -> List[str]:
    import pdfplumber

    pages: List[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return pages


def _ocr_pdf_pages(path: Path, page_indexes: List[int]) -> Dict[int, str]:
    try:
        import pypdfium2 as pdfium
        import pytesseract
    except Exception:
        return {}
    pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd
    if settings.tessdata_prefix:
        import os

        os.environ["TESSDATA_PREFIX"] = settings.tessdata_prefix

    results: Dict[int, str] = {}
    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception:
        return {}

    for idx in page_indexes:
        try:
            page = pdf[idx]
            image = page.render(scale=2).to_pil()
            text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            if text.strip():
                results[idx] = text
        except Exception:
            continue
    return results


def _read_text_with_quality(
    path: Path,
    enable_ocr: bool,
    min_ch_ratio: float,
    max_replacement_ratio: float,
    empty_page_threshold: float,
) -> Dict[str, object]:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
        return {"text": _normalize_from_pages([text]), "ocr_pages_processed": 0}
    if suffix == ".docx":
        import docx

        doc = docx.Document(str(path))
        raw = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return {"text": _normalize_from_pages([raw]), "ocr_pages_processed": 0}
    if suffix != ".pdf":
        raise ValueError(f"unsupported file type: {path.suffix}")

    pages = _extract_pdf_pages(path)
    empty_page_indexes = [i for i, p in enumerate(pages) if not p.strip()]
    low_quality_indexes = [
        i for i, p in enumerate(pages) if p.strip() and _is_low_quality(p, min_ch_ratio, max_replacement_ratio)
    ]

    should_ocr = False
    if pages:
        empty_ratio = len(empty_page_indexes) / len(pages)
        should_ocr = empty_ratio >= empty_page_threshold or len(low_quality_indexes) > 0

    ocr_pages_processed = 0
    if enable_ocr and should_ocr:
        target_indexes = sorted(set(empty_page_indexes + low_quality_indexes))
        ocr_map = _ocr_pdf_pages(path, target_indexes)
        for idx, txt in ocr_map.items():
            pages[idx] = txt
        ocr_pages_processed = len(ocr_map)

    normalized = _normalize_from_pages(pages)
    return {"text": normalized, "ocr_pages_processed": ocr_pages_processed}


class DocumentIngestor:
    def __init__(self, repository: ChromaPolicyRepository, index_path: str):
        self.repository = repository
        self.fingerprints = FingerprintStore(index_path=index_path)

    def ingest_path(
        self,
        docs_path: str,
        kb_scope: str,
        province_code: Optional[str],
        rebuild: bool,
        chunk_size: int,
        chunk_overlap: int,
        enable_ocr: bool,
        dedupe: bool,
        min_ch_ratio: float,
        max_replacement_ratio: float,
        empty_page_threshold: float,
    ) -> IngestStats:
        root = Path(docs_path)
        if not root.exists():
            raise FileNotFoundError(f"docs_path not found: {docs_path}")
        if not root.is_dir():
            raise ValueError(f"docs_path must be a directory: {docs_path}")
        files = [f for f in root.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED_SUFFIXES]

        stats = IngestStats(files_processed=len(files))
        all_chunks: List[str] = []
        all_metas: List[Dict[str, str]] = []

        index = self.fingerprints.load()
        if rebuild:
            index = self.fingerprints.clear_scope(index, kb_scope=kb_scope, province_code=province_code)

        for file in files:
            file_hash = _sha256(file)
            file_key = f"{kb_scope}:{(province_code or '').upper()}:{file.resolve()}"
            old_hash = index.get(file_key)

            if dedupe and not rebuild and old_hash == file_hash:
                stats.files_skipped += 1
                continue

            if dedupe and not rebuild and old_hash and old_hash != file_hash:
                self.repository.delete_by_file_hash(
                    kb_scope=kb_scope,
                    province_code=province_code,
                    file_hash=old_hash,
                )
                stats.files_updated += 1
            elif old_hash is None or rebuild:
                stats.files_new += 1

            extracted = _read_text_with_quality(
                file,
                enable_ocr=enable_ocr,
                min_ch_ratio=min_ch_ratio,
                max_replacement_ratio=max_replacement_ratio,
                empty_page_threshold=empty_page_threshold,
            )
            content = str(extracted.get("text") or "")
            stats.ocr_pages_processed += int(extracted.get("ocr_pages_processed") or 0)
            if not content.strip():
                continue

            chunks = split_text(content, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            title = file.stem
            effective_date = _extract_effective_date(content)
            policy_level = _extract_policy_level(content)
            for idx, chunk in enumerate(chunks):
                all_chunks.append(chunk)
                all_metas.append(
                    {
                        "province_code": province_code or "",
                        "doc_id": f"{file_hash}:{idx}",
                        "source_name": file.name,
                        "source_path": str(file.resolve()),
                        "file_hash": file_hash,
                        "doc_title": title,
                        "effective_date": effective_date,
                        "policy_level": policy_level,
                    }
                )
            index[file_key] = file_hash

        if all_chunks:
            self.repository.ingest_chunks(
                texts=all_chunks,
                metadatas=all_metas,
                kb_scope=kb_scope,
                province_code=province_code,
                rebuild=rebuild,
            )
            stats.chunks_created = len(all_chunks)

        self.fingerprints.save(index)
        return stats