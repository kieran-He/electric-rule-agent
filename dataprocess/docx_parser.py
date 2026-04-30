from __future__ import annotations

from pathlib import Path

from dataprocess.schemas import RawPage


class DocumentParseError(RuntimeError):
    pass


def parse_docx(file_path: str) -> list[RawPage]:
    path = Path(file_path)
    try:
        from docx import Document
    except Exception as exc:
        raise DocumentParseError(f"python-docx is unavailable: {exc}") from exc

    try:
        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        tables = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))
        text = "\n".join(paragraphs + tables)
        return [RawPage(page_number=1, text=text)]
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse DOCX {file_path}: {exc}") from exc